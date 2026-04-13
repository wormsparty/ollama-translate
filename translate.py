import requests
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL = "mistral-small3.2"

SYSTEM_PROMPT = (
    "Tu es un traducteur professionnel français vers allemand spécialisé dans les documents administratifs et appels d'offre suisses. "
    "Traduis le texte suivant du français vers l'allemand standard suisse (Schweizer Hochdeutsch) en respectant ces règles : "
    "utilise toujours 'ss' à la place de 'ß' (le ß n'existe pas en Suisse) ; "
    "utilise les guillemets suisses « » ; "
    "emploie le vocabulaire administratif helvétique (ex: 'Offerte' plutôt que 'Angebot', 'Submission', 'Kanton', 'Gemeinde', etc.) ; "
    "maintiens un registre formel et neutre adapté aux marchés publics suisses. "
    "Conserve la structure logique et les sauts de paragraphes. "
    "Ne rien ajouter ni retirer au contenu. "
    "Réponds uniquement avec la traduction, sans commentaire."
)


def iter_block_items(doc, parent=None):
    """Yield paragraphs and tables in document order, including inside w:sdt content controls."""
    if parent is None:
        parent = doc.element.body
    for child in parent.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)
        elif child.tag == qn("w:sdt"):
            sdt_content = child.find(qn("w:sdtContent"))
            if sdt_content is not None:
                yield from iter_block_items(doc, sdt_content)


def extract_header_text(doc) -> str:
    seen_texts = set()
    lines = []
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            if hdr is None:
                continue
            for p in hdr.paragraphs:
                t = p.text.strip()
                if t and t not in seen_texts:
                    seen_texts.add(t)
                    lines.append(t)
            for tbl in hdr.tables:
                for row in tbl.rows:
                    seen = set()
                    row_texts = []
                    for cell in row.cells:
                        if cell._tc not in seen:
                            seen.add(cell._tc)
                            if cell.text.strip():
                                row_texts.append(cell.text.strip())
                    if row_texts:
                        t = "\t".join(row_texts)
                        if t not in seen_texts:
                            seen_texts.add(t)
                            lines.append(t)
    return "\n".join(lines)


def extract_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    parts = []

    header_text = extract_header_text(doc)
    if header_text:
        parts.append(f"[EN-TÊTE]\n{header_text}\n[FIN EN-TÊTE]")

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if block.text.strip():
                parts.append(block.text)
        elif isinstance(block, Table):
            for row in block.rows:
                seen = set()
                row_texts = []
                for cell in row.cells:
                    if cell._tc not in seen:
                        seen.add(cell._tc)
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                if row_texts:
                    parts.append("\t".join(row_texts))

    return "\n\n".join(parts)


def translate_text(text: str) -> str:
    payload = {
        "model": MODEL,
        "system": SYSTEM_PROMPT,
        "prompt": text,
        "stream": False,
    }
    response = requests.post(OLLAMA_URL, json=payload, timeout=600)
    response.raise_for_status()
    return response.json()["response"].strip().replace("ß", "ss")


def translate_document(docx_path: Path, output_dir: Path) -> None:
    print(f"[→] Traitement : {docx_path.name}")

    text = extract_text(docx_path)
    if not text:
        print(f"  [!] Document vide, ignoré.")
        return

    word_count = len(text.split())
    print(f"  Mots détectés : {word_count}")
    print(f"  Traduction en cours...")

    translated = translate_text(text)

    output_path = output_dir / (docx_path.stem + ".md")
    output_path.write_text(translated, encoding="utf-8-sig")
    print(f"  [✓] Sauvegardé : {output_path}")


def main() -> None:
    input_dir = Path("fr")
    output_dir = Path("de")

    if not input_dir.exists():
        print(f"[!] Dossier source '{input_dir}' introuvable.")
        return

    output_dir.mkdir(exist_ok=True)

    docx_files = sorted(input_dir.glob("*.docx"))
    if not docx_files:
        print(f"[!] Aucun fichier .docx trouvé dans '{input_dir}'.")
        return

    print(f"[i] {len(docx_files)} fichier(s) trouvé(s) dans '{input_dir}'\n")

    for i, docx_path in enumerate(docx_files, 1):
        print(f"--- Fichier {i}/{len(docx_files)} ---")
        try:
            translate_document(docx_path, output_dir)
        except requests.exceptions.ConnectionError:
            print(f"  [✗] Impossible de joindre Ollama sur {OLLAMA_URL}. Est-il démarré ?")
        except requests.exceptions.Timeout:
            print(f"  [✗] Timeout pour {docx_path.name}.")
        except Exception as e:
            print(f"  [✗] Erreur inattendue : {e}")
        print()

    print("[i] Terminé.")


if __name__ == "__main__":
    main()